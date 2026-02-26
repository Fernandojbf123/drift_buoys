
from genericpath import exists

from docx.shared import Inches
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph

def insertar_parrafo_despues(paragraph, texto="", centrado=False):
    """Inserta un nuevo párrafo después del párrafo dado usando manipulación XML.
    
    Args:
        paragraph: El párrafo después del cual insertar.
        texto: El texto del nuevo párrafo (opcional).
        centrado: Si True, centra el texto del párrafo.
    
    Returns:
        El elemento XML del nuevo párrafo creado.
    """
    # Obtener el elemento del párrafo actual
    p_element = paragraph._element
    # Obtener el elemento padre
    parent = p_element.getparent()
    # Crear un nuevo elemento de párrafo
    nuevo_p = OxmlElement('w:p')
    
    # Si necesita estar centrado
    if centrado:
        pPr = OxmlElement('w:pPr')
        jc = OxmlElement('w:jc')
        jc.set(qn('w:val'), 'center')
        pPr.append(jc)
        nuevo_p.append(pPr)
    
    # Si hay texto, agregarlo
    if texto:
        run = OxmlElement('w:r')
        text_elem = OxmlElement('w:t')
        text_elem.text = texto
        run.append(text_elem)
        nuevo_p.append(run)
    
    # Insertar el nuevo párrafo después del actual
    parent.insert(parent.index(p_element) + 1, nuevo_p)
    
    return nuevo_p


def aux_insertar_figura_sin_titulo(paragraph, key, ruta_imagen, ancho_inches=6):
    """Inserta UNA imagen en un párrafo de Word SIN título (caso de array con un solo elemento).
    
    Args:
        paragraph: El párrafo del documento Word donde se buscará el marcador.
        key: El marcador de posición a buscar (por ejemplo, "<<fig_ejecucion_campania>>").
        ruta_imagen: La ruta del archivo de imagen a insertar.
        ancho_inches: El ancho de la imagen en pulgadas (por defecto 6).
    
    Returns:
        True si se insertó la imagen, False si no se encontró el marcador.
    """
    full_text = "".join(run.text for run in paragraph.runs)
    
    if key not in full_text:
        return False
    
    # Limpiar el párrafo (borrar todos los runs)
    for run in paragraph.runs:
        run.text = ""
    
    # Insertar la imagen en el párrafo
    run = paragraph.runs[0] if paragraph.runs else paragraph.add_run()
    run.add_picture(ruta_imagen, width=Inches(ancho_inches))
    
    # Insertar un salto de línea (párrafo vacío) después
    insertar_parrafo_despues(paragraph, texto="")
    
    return True


def aux_insertar_figuras_con_titulo(paragraph, key, lista_figuras):
    """Inserta MÚLTIPLES imágenes con sus títulos (caso de array con varios elementos).
    
    Args:
        paragraph: El párrafo del documento Word donde se buscará el marcador.
        key: El marcador de posición a buscar (por ejemplo, "<<fig_ejecucion_campania>>").
        lista_figuras: Lista de diccionarios con keys "ruta", "titulo", "tamanio".
    
    Returns:
        True si se insertaron las imágenes, False si no se encontró el marcador.
    """
    
    
    full_text = "".join(run.text for run in paragraph.runs)
    
    if key not in full_text:
        return False
    
    # Limpiar el párrafo (borrar todos los runs)
    for run in paragraph.runs:
        run.text = ""
    
    # Obtener referencias para inserción
    p_element = paragraph._element
    parent = p_element.getparent()
    indice_base = parent.index(p_element)
    
    # Insertar todas las figuras con sus títulos
    offset = 0
    for item in lista_figuras:
        ruta = item.get("ruta", "")
        titulo = item.get("titulo", "")
        ancho = item.get("tamanio", 6)
        
        # Si es la primera imagen, usar el párrafo actual
        if offset == 0:
            run = paragraph.runs[0] if paragraph.runs else paragraph.add_run()
            if exists(ruta):
                run.add_picture(ruta, width=Inches(ancho))
            offset = 1
        else:
            # Crear nuevo párrafo para imagen
            nuevo_p_img = OxmlElement('w:p')
            parent.insert(indice_base + offset, nuevo_p_img)
            # Convertir a Paragraph para poder agregar imagen
            para_img = Paragraph(nuevo_p_img, paragraph._parent)
            if exists(ruta):
                para_img.add_run().add_picture(ruta, width=Inches(ancho))
            offset += 1
        
        # Crear párrafo para el título (centrado)
        nuevo_p_titulo = OxmlElement('w:p')
        pPr = OxmlElement('w:pPr')
        jc = OxmlElement('w:jc')
        jc.set(qn('w:val'), 'center')
        pPr.append(jc)
        nuevo_p_titulo.append(pPr)
        run_titulo = OxmlElement('w:r')
        text_elem = OxmlElement('w:t')
        text_elem.text = titulo
        run_titulo.append(text_elem)
        nuevo_p_titulo.append(run_titulo)
        parent.insert(indice_base + offset, nuevo_p_titulo)
        offset += 1
        
        # Salto de línea
        nuevo_p_salto = OxmlElement('w:p')
        parent.insert(indice_base + offset, nuevo_p_salto)
        offset += 1
    
    return True


def aux_reemplazar_variable_en_parrafo(paragraph, key, value):
    """Reemplaza un marcador de posición en un párrafo de Word PRESERVANDO el formato.
    
    Args:
        paragraph: El párrafo donde buscar el marcador.
        key: El marcador de posición a buscar (por ejemplo, "<<orden_de_servicio>>").
        value: El texto que lo reemplazará (i.e., 100).
    
    Esta función preserva el formato del run donde EMPIEZA el marcador.
    Maneja tanto marcadores dentro de un solo run como divididos entre múltiples runs.
    """
    # Unir todo el texto del párrafo para verificar si este tiene al marcador
    full_text = "".join(run.text for run in paragraph.runs)
    if key not in full_text:
        return paragraph
    
    # Convertir value a string
    new_value = str(value[0]) if isinstance(value, list) else str(value)
    
    # CASO 1: Intentar el caso simple primero (todo en un run)
    for irun, run in enumerate(paragraph.runs):
        if key in run.text:
            run.text = run.text.replace(key, new_value)
        
        elif "<<" in run.text: # Si el marcador está dividido
            paragraph.runs[irun].text = new_value
            paragraph.runs[irun+1].text = ""
            paragraph.runs[irun+2].text = ""
            
            return paragraph
    
    
    
def reemplazar_en_word(doc, diccionario_de_reemplazos):
    """Reemplaza los marcadores de posición en un documento de Word utilizando un diccionario de reemplazos.
    doc es el documento de Word (objeto Document).
    diccionario_de_reemplazos es un diccionario donde las claves son los marcadores de posición a buscar
    (por ejemplo, "<<orden_de_servicio>>") y los valores son los textos que los reemplazarán (i.e., 100).
    
    Para las figuras, el valor debe ser una lista de diccionarios:
    - Lista con un solo elemento: inserta la figura SIN título
    - Lista con varios elementos: inserta las figuras CON sus títulos
    
    Cada diccionario debe tener las keys: "ruta", "titulo", "tamanio"
    """
    # Para cada párrafo en el documento, reemplaza los marcadores de posición utilizando el diccionario
    for variable, dato in diccionario_de_reemplazos.items():
        for parrafo in doc.paragraphs:
            if variable in parrafo.text:
                
                if "fig" not in variable: # Si el marcador no es de figura, reemplazo normal
                    aux_reemplazar_variable_en_parrafo(parrafo, variable, dato)
                
                # if "fig" in variable: # Si el marcador es de figura
                #     if isinstance(dato, list):
                #         if len(dato) == 1: # Una sola figura SIN título
                #             item = dato[0]
                #             ruta = item.get("ruta", "")
                #             ancho = item.get("tamanio", 6)
                #             aux_insertar_figura_sin_titulo(parrafo, variable, ruta, ancho)
                #         else: # Varias figuras CON títulos
                #             aux_insertar_figuras_con_titulo(parrafo, variable, dato)
                # else: # No es un marcador de figura, reemplazo normal
                #     aux_reemplazar_en_parrafo(parrafo, variable, dato)
    
    # for paragraph in doc.paragraphs:
        
    #     for key, value in diccionario_de_reemplazos.items():
    #         if "fig" not in key: # Si el marcador no es una figura, reemplaza normalmente
    #             aux_reemplazar_en_parrafo(paragraph, key, value)
            
    #         elif "fig" in key: # Si son las figuras
    #             if isinstance(value, list):
    #                 if len(value) == 1: # Si solo es una figura, insertarla SIN título
    #                     item = value[0]
    #                     ruta = item.get("ruta", "")
    #                     ancho = item.get("tamanio", 6)
    #                     aux_insertar_figura_sin_titulo(paragraph, key, ruta, ancho)
    #                 else: # Si son varias figuras, insertarlas CON títulos
    #                     aux_insertar_figuras_con_titulo(paragraph, key, value)
    
    # Retornar el documento modificado
    return doc
            
    # # Ahora en las tablas
    # for table in doc.tables:
    #     for row in table.rows:
    #         for cell in row.cells:
    #             for paragraph in cell.paragraphs:
    #                 for key, value in diccionario_de_reemplazos.items():
    #                     if "fig" not in key: # No es una figura
    #                         aux_reemplazar_en_parrafo(paragraph, key, value)
    #                     else: # Es una figura
    #                         if isinstance(value, list):
    #                             if len(value) == 1: # Una sola figura SIN título
    #                                 item = value[0]
    #                                 ruta = item.get("ruta", "")
    #                                 ancho = item.get("tamanio", 6)
    #                                 aux_insertar_figura_sin_titulo(paragraph, key, ruta, ancho)
    #                             else: # Varias figuras CON títulos
    #                                 aux_insertar_figuras_con_titulo(paragraph, key, value)